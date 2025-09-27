import streamlit as st
import pandas as pd
import yfinance as yf
import numpy as np
import plotly.graph_objects as go
import os
import requests
from bs4 import BeautifulSoup
import re
import json # Added for NSE API processing
import time # FIX: Added to resolve NameError for time.sleep()
from typing import Dict, Any, List # Added for typing in helper functions
from io import BytesIO # For document download
from docx import Document # Added for Word document generation
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ta.trend import MACD
from ta.momentum import RSIIndicator
from ta.volatility import BollingerBands, AverageTrueRange
from ta.volume import OnBalanceVolumeIndicator

# ======================================================================================
# CONFIGURATION & HEADER
# ======================================================================================
st.set_page_config(page_title="Stock Analysis Dashboard", layout="wide")

st.title("Stock Analyzer | NS   T R A D E R")
st.markdown("Select an industry from your Excel file to get a consolidated analysis, including financial ratios from **Screener.in** and a detailed **Swing Trading** recommendation with Fibonacci levels.")

# --- NSE Configuration (Used for data export) ---
# NOTE: The ISIN here is hardcoded to INFY (INE009A01021) as the Streamlit app does not track ISINs.
# The export will only work correctly for INFY's ownership data unless you map ISINs in your Excel file.
NSE_BASE_URL = "https://www.nseindia.com/"

HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/555.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/555.36',
    'Accept-Language': 'en-US,en;q=0.9',
    'Accept-Encoding': 'gzip, deflate, br',
}
# --- End NSE Configuration ---

# ======================================================================================
# DATA FETCHING & CALCULATION FUNCTIONS (Dashboard UI)
# ======================================================================================

@st.cache_data(ttl=3600) # Cache data for 1 hour
def get_stock_data(ticker):
    """Fetches all necessary data for a stock from yfinance."""
    stock = yf.Ticker(f"{ticker}.NS")
    history_weekly = stock.history(period="3y", interval="1wk")
    history_daily = stock.history(period="2y", interval="1d")
    info = stock.info
    financials = stock.financials
    return history_weekly, info, financials, history_daily

@st.cache_data(ttl=3600)
def get_price_on_date(daily_history, target_date_str):
    """Finds the closest closing price and the actual date to a target date from daily history."""
    try:
        target_date = pd.to_datetime(target_date_str)
        closest_date_index = daily_history.index.get_indexer([target_date], method='nearest')[0]
        actual_date = daily_history.index[closest_date_index]
        price = daily_history.iloc[closest_date_index]['Close']
        return price, actual_date
    except Exception:
        return None, None

@st.cache_data(ttl=3600)
def scrape_screener_data(ticker):
    """Scrapes key financial data for a given ticker from screener.in."""
    url = f"https://www.screener.in/company/{ticker}/consolidated/"
    try:
        response = requests.get(url, headers=HEADERS, timeout=10)
        response.raise_for_status()
    except requests.RequestException:
        return None, "Failed to load page"
    
    soup = BeautifulSoup(response.content, 'html.parser')
    data = {}
    ratio_list = soup.select_one('#top-ratios')
    if not ratio_list: return None, "Ratios not found"
    for li in ratio_list.find_all('li'):
        name = li.select_one('.name').get_text(strip=True) if li.select_one('.name') else ''
        value = li.select_one('.nowrap.value .number').get_text(strip=True) if li.select_one('.nowrap.value .number') else ''
        if name and value: data[name] = value
    return data, "Success"

def calculate_graham_intrinsic_value(info, financials, bond_yield=7.5):
    """Calculates the intrinsic value of a stock using Benjamin Graham's formula."""
    try:
        eps = info.get('trailingEps')
        if not eps or eps <= 0: return None
        net_income = financials.loc['Net Income']
        if net_income.isnull().all() or len(net_income.dropna()) < 2: return None
        growth_rates = net_income.pct_change().dropna()
        avg_growth_rate = np.mean(growth_rates)
        g = min(avg_growth_rate * 100, 15.0)  
        if g < 0: g = 0
        return (eps * (8.5 + 2 * g) * 4.4) / bond_yield
    except (KeyError, IndexError, TypeError):
        return None

# --- FIBONACCI RETRACEMENT ANALYSIS ---
def calculate_fibonacci_levels(history):
    # ... (function body remains the same as previous code)
    lookback_period = history.tail(52) # Look at the last 52 weeks (1 year)
    high_price = lookback_period['High'].max()
    low_price = lookback_period['Low'].min()
    price_range = high_price - low_price
    current_price = history['Close'].iloc[-1]

    # Determine trend
    is_uptrend = current_price > lookback_period['Close'].iloc[0]

    levels = {}
    if is_uptrend:
        levels['23.6%'] = high_price - (price_range * 0.236)
        levels['38.2%'] = high_price - (price_range * 0.382)
        levels['50.0%'] = high_price - (price_range * 0.500)
        levels['61.8%'] = high_price - (price_range * 0.618)
    else: # Downtrend
        levels['23.6%'] = low_price + (price_range * 0.236)
        levels['38.2%'] = low_price + (price_range * 0.382)
        levels['50.0%'] = low_price + (price_range * 0.500)
        levels['61.8%'] = low_price + (price_range * 0.618)
    
    # Generate signal
    signal = "Neutral"
    if is_uptrend:
        if current_price > levels['38.2%'] and current_price < high_price:
            signal = f"Finding support above the 38.2% level (₹{levels['38.2%']:.2f}). Potential continuation of uptrend."
        elif current_price <= levels['61.8%']:
            signal = "Trend weakening, has broken below the 61.8% support."
    else: # Downtrend
        if current_price < levels['61.8%'] and current_price > low_price:
            signal = f"Facing resistance below the 61.8% level (₹{levels['61.8%']:.2f}). Potential continuation of downtrend."
        elif current_price >= levels['61.8%']:
            signal = "Potential trend reversal, has broken above the 61.8% resistance."
            
    return levels, signal, is_uptrend, high_price, low_price


def calculate_swing_trade_analysis(history):
    # ... (function body remains the same as previous code)
    if len(history) < 52:
        return None, "Insufficient Data", "Not enough weekly data for full analysis."

    close = history['Close']
    price = close.iloc[-1]
    
    sma_20 = close.rolling(window=20).mean().iloc[-1]
    sma_50 = close.rolling(window=50).mean().iloc[-1]
    rsi_14 = RSIIndicator(close, window=14).rsi().iloc[-1]
    macd_indicator = MACD(close, window_slow=26, window_fast=12, window_sign=9)
    macd_line = macd_indicator.macd().iloc[-1]
    macd_signal = macd_indicator.macd_signal().iloc[-1]
    macd_hist = macd_indicator.macd_diff().iloc[-1]
    bb_indicator = BollingerBands(close, window=20, window_dev=2)
    bb_high = bb_indicator.bollinger_hband().iloc[-1]
    bb_low = bb_indicator.bollinger_lband().iloc[-1]
    obv_indicator = OnBalanceVolumeIndicator(close, history['Volume'])
    obv_slope = obv_indicator.on_balance_volume().diff().rolling(window=5).mean().iloc[-1]
    atr_14 = AverageTrueRange(history['High'], history['Low'], close, window=14).average_true_range().iloc[-1]

    indicators = {"20W SMA": sma_20, "50W SMA": sma_50, "RSI (14)": rsi_14, "MACD Line": macd_line, "MACD Signal": macd_signal}

    score = 0
    reasons = []
    
    if price > sma_20: score += 2; reasons.append("✅ Price > 20W SMA (Short-term trend is up).")
    else: reasons.append("❌ Price < 20W SMA (Short-term trend is down).")
    if sma_20 > sma_50: score += 2; reasons.append("✅ 20W SMA > 50W SMA (Golden Cross).")
    else: reasons.append("❌ 20W SMA < 50W SMA (Death Cross).")
    if macd_line > macd_signal: score += 1; reasons.append("✅ MACD > Signal (Bullish momentum).")
    else: reasons.append("❌ MACD < Signal (Bearish momentum).")
    if 45 < rsi_14 < 68: score += 2; reasons.append(f"✅ RSI is healthy at {rsi_14:.1f}.")
    else: reasons.append(f"⚠️ RSI is {rsi_14:.1f} (Not in optimal range).")
    if obv_slope > 0: score += 2; reasons.append("✅ OBV trend is positive (Volume confirms trend).")
    else: reasons.append("❌ OBV trend is negative (Volume does not confirm).")
    
    if score >= 7: recommendation = "Strong Buy"
    elif score >= 5: recommendation = "Buy"
    elif score >= 3: recommendation = "Hold / Monitor"
    else: recommendation = "Sell / Avoid"
    
    return indicators, recommendation, "\n\n".join(reasons)


@st.cache_data(ttl=3600)
def calculate_quick_signals(df, ticker_col):
    # ... (function body remains the same as previous code)
    tickers = df[ticker_col].dropna().unique()
    all_signals = []
    for ticker in tickers[:50]: 
        try:
            history, _, _, _ = get_stock_data(ticker)
            if not history.empty and len(history) >= 52:
                _, recommendation, _ = calculate_swing_trade_analysis(history)
                all_signals.append({'Ticker': ticker, 'Signal': recommendation})
        except Exception: continue
    if not all_signals: return pd.DataFrame(), pd.DataFrame()
    signals_df = pd.DataFrame(all_signals)
    buy_signals = signals_df[signals_df['Signal'].isin(['Strong Buy', 'Buy'])].head(3)
    sell_signals = signals_df[signals_df['Signal'] == 'Sell / Avoid'].head(3)
    return buy_signals, sell_signals

# ======================================================================================
# WORD DOCUMENT GENERATION MODULE (Integrated from stock_report_generator.py)
# ======================================================================================

def create_nse_session() -> requests.Session:
    """Creates a session and fetches initial cookies needed for NSE API access."""
    session = requests.Session()
    try:
        session.get(NSE_BASE_URL, headers=HEADERS, timeout=10)
        return session
    except requests.RequestException:
        return None

def fetch_nse_quote_data_export(session: requests.Session, ticker: str) -> Dict[str, Any]:
    """
    Fetches the main stock quote data from the NSE API, including all detailed trade metrics.
    """
    data = {}
    NSE_QUOTE_API = f"https://www.nseindia.com/api/quote-equity?symbol={ticker}"
    try:
        response = session.get(NSE_QUOTE_API, headers=HEADERS, timeout=10)
        response.raise_for_status()
        quote_data = response.json()
        
        # --- Extract metric groups ---
        info = quote_data.get('info', {})
        metadata = quote_data.get('metadata', {})
        security_info = quote_data.get('securityInfo', {})
        price_band = quote_data.get('priceBand', {})
        p_data = quote_data.get('preOpenMarket', {}).get('finalPrice', {})
        
        # --- Calculations and Conversions ---
        volume = quote_data.get('totalTradedVolume', 0)
        value = quote_data.get('totalTradedValue', 0)
        delivery_qty = security_info.get('deliveryQuantity', 0)
        
        deliverable_percentage = (delivery_qty / volume * 100) if volume > 0 else 0
        
        # Use day high/low for indicative daily volatility (as a percentage of day average)
        day_high = quote_data.get('dayHigh', 0)
        day_low = quote_data.get('dayLow', 0)
        daily_volatility_percent = (day_high - day_low) / day_low * 100 if day_low > 0 else 0
        
        # --- Final Data Structure for Export ---
        data['Latest Price Data'] = {
            'Company Name': info.get('companyName', 'N/A'),
            'Basic Industry': info.get('industry', 'N/A'),
            'Latest Trade Date': metadata.get('lastUpdateTime', 'N/A'),
            'Total Capital Market (₹ Crore)': f"{round(info.get('marketCap', 0) / 10000000, 2):,}",
            
            # Trade Volume/Value
            'Total Traded Volume (Lakhs)': f"{round(volume / 100000, 2):,}",
            'Total Traded Value (₹ Crores)': f"{round(value / 10000000, 2):,}",
            'Delivery Percentage': f"{deliverable_percentage:.2f}%",

            # Price Metrics
            'Face Value (₹)': security_info.get('faceValue', 'N/A'),
            'Date of Listing': security_info.get('listingDate', 'N/A'),
            
            '52 Week High (₹)': metadata.get('weekHigh', 'N/A'),
            '52 Week High Date': metadata.get('weekHighDate', 'N/A'),
            '52 Week Low (₹)': metadata.get('weekLow', 'N/A'),
            '52 Week Low Date': metadata.get('weekLowDate', 'N/A'),
            
            'Upper Circuit Band (₹)': price_band.get('max', 'N/A'),
            'Lower Circuit Band (₹)': price_band.get('min', 'N/A'),
            
            # Volatility & P/E Ratios
            'Daily Volatility (Indicative)': f"{daily_volatility_percent:.2f}%",
            'Annualised Volatility': 'N/A (API Field Missing)', # Not directly available in quote API
            'Adjusted P/E (TTM)': p_data.get('pE', 'N/A'),
            'Symbol P/E': metadata.get('symbolPE', 'N/A'),
        }
    except Exception:
        # Catch any errors during API call or processing
        pass
    return data

def fetch_nse_financial_data_export(session: requests.Session, ticker: str) -> pd.DataFrame:
    """Fetches quarterly financial results for export."""
    df = pd.DataFrame()
    NSE_FINANCIAL_API = f"https://www.nseindia.com/api/corporates-financial-results?symbol={ticker}&index=equities&period=Quarterly"
    time.sleep(1) 
    
    try:
        response = session.get(NSE_FINANCIAL_API, headers=HEADERS, timeout=10)
        response.raise_for_status()
        financial_data = response.json()
        
        if 'data' in financial_data:
            results = []
            for item in financial_data['data']:
                income = item.get('totalIncome', 0)
                expenses = item.get('totalExpenses', 0)
                tax = item.get('taxExpense', 0)
                
                results.append({
                    'Quarter Ended': item.get('period'),
                    'Total Income (Cr)': round(income / 100, 2), 
                    'Total Expenses (Cr)': round(expenses / 100, 2), 
                    'Total Tax Expense (Cr)': round(tax / 100, 2)
                })
            
            df = pd.DataFrame(results).head(4)
            if not df.empty:
                df = df.set_index('Quarter Ended').T
            
    except Exception:
        pass
    return df

def fetch_nse_shareholding_data_export(session: requests.Session, ticker: str) -> pd.DataFrame:
    """Fetches the latest quarterly FII/DII shareholding pattern for export."""
    df = pd.DataFrame()
    NSE_SHAREHOLDING_API = f"https://www.nseindia.com/api/corporates-shareholding?symbol={ticker}"
    time.sleep(1)
    
    try:
        response = session.get(NSE_SHAREHOLDING_API, headers=HEADERS, timeout=10)
        response.raise_for_status()
        shareholding_data = response.json()
        
        if 'data' in shareholding_data:
            results = []
            for item in shareholding_data['data']:
                category = item.get('category', 'N/A')
                percent = item.get('value', '0.0')
                
                if 'FII' in category.upper() or 'DII' in category.upper() or 'MUTUAL FUND' in category.upper():
                    results.append({'Category': category, 'Percentage (%)': float(percent)})
            
            latest_date = shareholding_data.get('latest_date', 'N/A')
            
            if results:
                df = pd.DataFrame(results)
                df_agg = df.groupby('Category')['Percentage (%)'].sum().reset_index()
                df_agg['Percentage (%)'] = df_agg['Percentage (%)'].apply(lambda x: f"{x:.2f}%")
                df_agg.columns = ['Category', f'Latest Percentage ({latest_date})']
                df = df_agg.set_index('Category').T
        
    except Exception:
        pass
    return df

def fetch_cogencis_ownership_data_export(isin: str, ticker: str) -> Dict[str, pd.DataFrame]:
    """Scrapes all tables from the Cogencis ownership data page for export."""
    scraped_data = {}
    # NOTE: The ISIN here is hardcoded to INFY (INE009A01021) 
    # and the ticker is dynamically inserted into the URL.
    COGENCIS_OWNERSHIP_URL = f"https://iinvest.cogencis.com/{isin}/symbol/ns/{ticker}/Infosys%20Limited?tab=ownership-data&type=capital-history"
    time.sleep(2)
    
    try:
        response = requests.get(COGENCIS_OWNERSHIP_URL, headers=HEADERS, timeout=15)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        tables = soup.find_all('table')
        
        for i, table in enumerate(tables):
            try:
                df_list = pd.read_html(str(table))
                if df_list:
                    df = df_list[0]
                    title_tag = table.find_previous(['h3', 'h4', 'h5', 'p'], text=True)
                    table_title = title_tag.text.strip() if title_tag and len(title_tag.text.strip()) > 5 else f"Ownership/Capital Table {i+1}"
                    
                    if not df.empty:
                        if isinstance(df.columns, pd.MultiIndex):
                            df.columns = [' '.join(col).strip() for col in df.columns.values]
                        scraped_data[table_title] = df.fillna('')
            except Exception:
                continue

    except Exception as e:
        scraped_data['Error'] = f"Failed to fetch Cogencis data: {e}"
        
    return scraped_data

def add_dataframe_to_word(document, df: pd.DataFrame, table_style: str = 'Table Grid'):
    """Helper function to convert a Pandas DataFrame to a Word table."""
    document.add_paragraph()
    has_index = df.index.name is not None or not pd.RangeIndex(start=0, stop=len(df.index)).equals(df.index)

    rows, cols = df.shape
    num_cols = cols + (1 if has_index else 0)
    table = document.add_table(rows + 1, num_cols)
    table.style = table_style
    
    hdr_cells = table.rows[0].cells
    current_col = 0
    if has_index:
        hdr_cells[0].text = str(df.index.name or 'Index')
        current_col = 1
        
    for j, col_name in enumerate(df.columns):
        hdr_cells[current_col + j].text = str(col_name)
        hdr_cells[current_col + j].paragraphs[0].runs[0].font.bold = True
        hdr_cells[current_col + j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i, (index_name, row) in enumerate(df.iterrows()):
        row_cells = table.rows[i + 1].cells
        current_col = 0
        
        if has_index:
            row_cells[0].text = str(index_name)
            current_col = 1
            
        for j, value in enumerate(row):
            text_value = str(value)
            try:
                if text_value.replace(',', '').replace('.', '', 1).isdigit() and len(text_value.replace('.', '')) > 4:
                    text_value = f"{float(text_value.replace(',', '')):,}"
            except:
                pass

            row_cells[current_col + j].text = text_value
            if any(char.isdigit() for char in text_value) and not any(char.isalpha() for char in text_value):
                 row_cells[current_col + j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                 row_cells[current_col + j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

def generate_word_document_bytes(ticker: str, isin: str = "INE009A01021") -> BytesIO:
    """Fetches all data, generates the Word document, and returns it as a BytesIO object."""
    nse_session = create_nse_session()
    
    if not nse_session:
        # Return an empty buffer if session fails
        return BytesIO(b"")

    # Fetch Data
    quote_data = fetch_nse_quote_data_export(nse_session, ticker)
    financial_data = fetch_nse_financial_data_export(nse_session, ticker)
    shareholding_data = fetch_nse_shareholding_data_export(nse_session, ticker)
    ownership_data = fetch_cogencis_ownership_data_export(isin, ticker)

    document = Document()
    
    # Set document style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Main Title
    document.add_heading(f"Comprehensive Stock Analysis Report: {ticker}", 0)
    document.add_paragraph(f"Report Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph("---")

    # --- Section 1: Trade Information and Price Metrics (NSE) ---
    document.add_heading(f"1. Trade Information and Price Metrics (NSE)", level=1)
    if 'Latest Price Data' in quote_data:
        data = quote_data['Latest Price Data']
        
        # Structure the trade data into a key-value DataFrame for clean presentation
        trade_info_list = [
            ("Company Name", data.get('Company Name', 'N/A')),
            ("Basic Industry", data.get('Basic Industry', 'N/A')),
            ("Last Updated", data.get('Latest Trade Date', 'N/A')),
            ("", "---"), 
            ("Total Capital Market", f"₹{data.get('Total Capital Market (₹ Crore)', 'N/A')} Crore"),
            ("Date of Listing", data.get('Date of Listing', 'N/A')),
            ("", "---"), 
            ("Total Traded Volume", f"{data.get('Total Traded Volume (Lakhs)', 'N/A')} Lakhs"),
            ("Total Traded Value", f"₹{data.get('Total Traded Value (₹ Crores)', 'N/A')} Crores"),
            ("% of Deliverable / Traded Qty", data.get('Delivery Percentage', 'N/A')),
            ("", "---"),
            ("Adjusted P/E (TTM)", data.get('Adjusted P/E (TTM)', 'N/A')),
            ("Symbol P/E", data.get('Symbol P/E', 'N/A')),
            ("Face Value", f"₹{data.get('Face Value (₹)', 'N/A')}"),
            ("", "---"),
            ("52 Week High Amount", f"₹{data.get('52 Week High (₹)', 'N/A')}"),
            ("52 Week High Date", data.get('52 Week High Date', 'N/A')),
            ("52 Week Low Amount", f"₹{data.get('52 Week Low (₹)', 'N/A')}"),
            ("52 Week Low Date", data.get('52 Week Low Date', 'N/A')),
            ("", "---"),
            ("Upper Circuit Band", f"₹{data.get('Upper Circuit Band (₹)', 'N/A')}"),
            ("Lower Circuit Band", f"₹{data.get('Lower Circuit Band (₹)', 'N/A')}"),
            ("Daily Volatility (Indicative)", data.get('Daily Volatility (Indicative)', 'N/A')),
            ("Annualised Volatility", data.get('Annualised Volatility', 'N/A')),
        ]
        
        df_trade_info = pd.DataFrame(trade_info_list, columns=['Metric', 'Value'])
        
        # Prepare for Word table (transpose to vertical list)
        df_trade_info_filtered = df_trade_info[df_trade_info['Metric'] != ''].set_index('Metric').T

        # Add the detailed table
        add_dataframe_to_word(document, df_trade_info_filtered, table_style='Table Grid')

    else:
        document.add_paragraph("Trade information could not be retrieved from NSE.")
    
    document.add_paragraph("---")

    # --- Section 2: Quarterly Financial Results Comparison (NSE) ---
    document.add_heading("2. Quarterly Financial Results Comparison (₹ Crores)", level=1)
    if not financial_data.empty:
        add_dataframe_to_word(document, financial_data, table_style='Light Grid Accent 2')
    else:
        document.add_paragraph("Quarterly financial comparison data could not be retrieved from NSE.")

    document.add_paragraph("---")

    # --- Section 3: FII/DII Shareholding Pattern (NSE) ---
    document.add_heading("3. Institutional Shareholding Pattern (NSE)", level=1)
    if not shareholding_data.empty:
        add_dataframe_to_word(document, shareholding_data, table_style='Grid Table 4 Accent 1')
        document.add_paragraph("Data represents the latest quarterly percentage breakdown reported by the company.")
    else:
        document.add_paragraph("Quarterly shareholding data (FII/DII breakdown) could not be retrieved from NSE.")
    
    document.add_paragraph("---")

    # --- Section 4: Cogencis Ownership Data ---
    document.add_heading(f"4. Ownership and Capital History (Source: Cogencis)", level=1)

    if 'Error' in ownership_data:
        document.add_paragraph(f"Error scraping Cogencis data: {ownership_data['Error']}")
    elif ownership_data:
        for title, df in ownership_data.items():
            document.add_heading(f"4.{list(ownership_data.keys()).index(title) + 1}: {title}", level=2)
            add_dataframe_to_word(document, df, table_style='List Table 4 Accent 3')
    else:
        document.add_paragraph("No ownership or capital history tables were successfully scraped from Cogencis.")
        
    # Save document to a BytesIO stream
    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io

# ======================================================================================
# STREAMLIT UI & LOGIC (Updated display_stock_analysis)
# ======================================================================================

def display_stock_analysis(ticker):
    """Displays analysis for a single stock."""
    try:
        history, info, financials, daily_history = get_stock_data(ticker)
        if history.empty:
            st.warning(f"Could not fetch price history for **{ticker}**. Skipping.")
            return

        st.header(f"Analysis for: {info.get('shortName', ticker)} ({ticker})", divider='rainbow')

        swing_indicators, swing_recommendation, _ = calculate_swing_trade_analysis(history)
        intrinsic_value = calculate_graham_intrinsic_value(info, financials)
        
        current_price = history['Close'].iloc[-1]
        price_mar_28, date_mar_28 = get_price_on_date(daily_history, '2025-03-28')
        move_fy_percent = None
        fy_delta_text = "N/A"
        if price_mar_28 and current_price:
            move_fy_percent = ((current_price - price_mar_28) / price_mar_28) * 100
            fy_delta_text = f"from ₹{price_mar_28:,.2f} on {date_mar_28.strftime('%d-%b-%Y')}"

        m_col1, m_col2, m_col3, m_col4, m_col5 = st.columns(5)
        m_col1.metric("Current Price", f"₹{current_price:,.2f}")
        m_col2.metric("Swing Signal", swing_recommendation)
        m_col3.metric("Intrinsic Value", f"₹{intrinsic_value:,.2f}" if intrinsic_value else "N/A")
        m_col4.metric(label="Move within FY", value=f"{move_fy_percent:.2f}%" if move_fy_percent is not None else "N/A", delta=fy_delta_text)
        m_col5.metric("Buy Price (≈20W SMA)", f"₹{swing_indicators['20W SMA']:,.2f}" if swing_indicators else "N/A")

        st.divider()

        chart_col, analysis_col = st.columns([2, 1])

        with chart_col:
            st.subheader("Weekly Price Chart with Fibonacci Retracement")
            fib_levels, _, is_uptrend, high_price, low_price = calculate_fibonacci_levels(history)
            
            fig = go.Figure(data=[go.Candlestick(x=history.index, open=history['Open'], high=history['High'], low=history['Low'], close=history['Close'], name='Price')])
            history['SMA_20W'] = history['Close'].rolling(window=20).mean()
            history['SMA_50W'] = history['Close'].rolling(window=50).mean()
            fig.add_trace(go.Scatter(x=history.index, y=history['SMA_20W'], mode='lines', name='20W SMA', line=dict(color='orange', width=1.5)))
            fig.add_trace(go.Scatter(x=history.index, y=history['SMA_50W'], mode='lines', name='50W SMA', line=dict(color='purple', width=1.5)))
            
            # Add Fibonacci lines to chart
            colors = ['red', 'orange', 'yellow', 'green']
            for i, (level, price) in enumerate(fib_levels.items()):
                fig.add_hline(y=price, line_width=1, line_dash="dash", line_color=colors[i], annotation_text=f"Fib {level}", annotation_position="bottom right")

            fig.update_layout(height=600, yaxis_title='Price (INR)', xaxis_rangeslider_visible=False, legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1))
            st.plotly_chart(fig, use_container_width=True)
            
        with analysis_col:
            st.subheader("Swing Signal Reasoning")
            _, _, swing_reasoning = calculate_swing_trade_analysis(history)
            st.info(swing_reasoning)

            # --- NEW: Fibonacci section ---
            st.subheader("Fibonacci Retracement Analysis")
            _, fib_signal, _, _, _ = calculate_fibonacci_levels(history)
            st.info(fib_signal)

            st.subheader("Key Financial Ratios")
            screener_data, screener_status = scrape_screener_data(ticker)
            if screener_status == "Success":
                df_screener = pd.DataFrame(screener_data.items(), columns=['Ratio', 'Value'])
                st.dataframe(df_screener, use_container_width=True, hide_index=True)
            else:
                st.warning(f"Could not scrape data ({screener_status}).")
        
        # --- Export Button Section (Placed below the columns) ---
        st.divider()
        
        export_col, _ = st.columns([1.5, 5.5])
        with export_col:
            # Generate the document bytes dynamically
            doc_bytes = generate_word_document_bytes(ticker)
            
            st.download_button(
                label="⬇️ Export Full Report (DOCX)",
                data=doc_bytes,
                file_name=f"{ticker}_Full_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
        # --- End Export Section ---


    except Exception as e:
        st.error(f"An error occurred while processing **{ticker}**: {e}")

# ======================================================================================
# STREAMLIT UI & LOGIC
# ======================================================================================
EXCEL_FILE_PATH = "SELECTED STOCKS 22FEB2025.xlsx"
TICKER_COLUMN_NAME = "NSE SYMBOL"
INDUSTRY_COLUMN_NAME = "INDUSTRY"

if 'current_stock_index' not in st.session_state: st.session_state.current_stock_index = 0
if 'ticker_list' not in st.session_state: st.session_state.ticker_list = []
if 'quick_signals_calculated' not in st.session_state: st.session_state.quick_signals_calculated = False

if not os.path.exists(EXCEL_FILE_PATH):
    st.error(f"Error: The file '{EXCEL_FILE_PATH}' was not found.")
else:
    with st.sidebar:
        st.header("⚙️ Analysis Filter")
        try:
            df_full = pd.read_excel(EXCEL_FILE_PATH, sheet_name='Sheet1')
            industries = ["All Industries"] + sorted(df_full[INDUSTRY_COLUMN_NAME].dropna().unique().tolist())
            selected_industry = st.selectbox("Select an Industry:", industries)
            
            if st.button("🚀 Analyze Selected Industry", type="primary"):
                df_filtered = df_full[df_full[INDUSTRY_COLUMN_NAME] == selected_industry] if selected_industry != "All Industries" else df_full
                st.session_state.ticker_list = df_filtered[TICKER_COLUMN_NAME].dropna().unique().tolist()
                st.session_state.current_stock_index = 0
                st.session_state.quick_signals_calculated = True
                
            if st.session_state.quick_signals_calculated:
                with st.spinner("Calculating market snapshot..."):
                    buy_signals, sell_signals = calculate_quick_signals(df_full, TICKER_COLUMN_NAME)
                
                st.subheader("Quick Signals Snapshot", divider='rainbow')
                st.markdown("**Top 3 Buy Signals**")
                if not buy_signals.empty:
                    for _, row in buy_signals.iterrows(): st.success(f"**{row['Ticker']}**: {row['Signal']}")
                else: st.info("No strong buy signals found.")
                
                st.markdown("**Top 3 Sell Signals**")
                if not sell_signals.empty:
                    for _, row in sell_signals.iterrows(): st.error(f"**{row['Ticker']}**: {row['Signal']}")
                else: st.info("No strong sell signals found.")
        except Exception as e:
            st.error(f"Could not read the Excel file. Error: {e}")

    if st.session_state.ticker_list:
        current_ticker = st.session_state.ticker_list[st.session_state.current_stock_index]
        
        col1, col2, col3 = st.columns([1.5, 5, 1.5])
        with col1:
            if st.button("⬅️ Previous Stock", use_container_width=True, disabled=(st.session_state.current_stock_index == 0)):
                st.session_state.current_stock_index -= 1; st.rerun()
        with col2:
            st.markdown(f"<p style='text-align: center; font-size: 1.1em;'>Displaying <b>{st.session_state.current_stock_index + 1}</b> of <b>{len(st.session_state.ticker_list)}</b> stocks</p>", unsafe_allow_html=True)
        with col3:
            if st.button("Next Stock ➡️", use_container_width=True, disabled=(st.session_state.current_stock_index >= len(st.session_state.ticker_list) - 1)):
                st.session_state.current_stock_index += 1; st.rerun()

        display_stock_analysis(current_ticker)
    else:
        st.info("Select an industry from the sidebar and click the 'Analyze' button to begin.")
